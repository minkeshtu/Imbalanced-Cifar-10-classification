import os
import glob
import numpy as np
import torch
from docx import Document
from docx.shared import Inches
import matplotlib
import matplotlib.patches as patches
import matplotlib.pyplot as plt
from sklearn.manifold import TSNE
from dataset.cifar10Loader import loadCIFAR10_testset, to_numpy
from architectures import Unet_based, Mobilenetv2_based

def plot_tsne_clusters(embedded_features, labels, num_of_images, label_map=None, vector_dim=512, doc_file=None, show=False):
    ''' Plot the encoded features clusters into 2-dementional space using t-sne
    params:
        embedded_features: python list of feature vector arrays
        labels: python list of labels correponds to the embedded_features
        number of images: int, number of images to plot
        label_map: python dict, cifar10 labels with thier corresponding index
        vector_dim: int, feature vector dims
        doc_file: str, docx file path to save plot
        show: bool, to visualize the clusters separately
    '''
    if label_map is None:
        label_map = {0:'airplane',1:'automobile',2:'bird',3:'cat',4:'deer',5:'dog',6:'frog',7:'horse',8:'ship',9:'truck'}
    fig = plt.figure(figsize=(10,10))
    fig, ax = plt.subplots()
    X_embedded = TSNE(n_components=2).fit_transform(embedded_features)
    count = 0
    for i, t in enumerate(sorted(set(labels))):
        idx = labels == t
        ax.scatter(X_embedded[idx, 0], X_embedded[idx, 1], label=label_map[t])
        count += 1
    leg = ax.legend(bbox_to_anchor=(1, 1))
    if doc_file:
        tmp_path = 'temp.jpg'
        plt.savefig(tmp_path, bbox_inches='tight')
        if os.path.exists(doc_file):
            document = Document(doc_file)
        else:
            document = Document()
        sep = '************************************************************'
        p = document.add_paragraph(sep)
        r = p.add_run(f'\nClusters of encoded features of {num_of_images} images:')
        r.bold = True
        document.add_picture(tmp_path, width=Inches(5))
        document.save(doc_file)
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    if show:
        plt.show()


def get_features(args, model_path, vector_dim, num_of_images=1000):
    ''' Get encoded features (encoder output) from model on testdataset [:num_of_images)
    '''
    filters =  filters = [int(args.width_scale_factor*i) for i in [32, 64, 128, 256]]
    if (model_path.find('unet_based') != -1):
        encoder = Unet_based.Encoder_for_visualization(3, filters)
    else:
        encoder = Mobilenetv2_based.Encoder_for_visualization()
    encoder.load_state_dict(torch.load(model_path), strict=False)
    encoder.eval()
    print('%s weights loaded successfully!'%model_path)

    testLoader = loadCIFAR10_testset(batch_size = num_of_images)
    print('%s images to embed'%num_of_images)
    for images, labels in testLoader:
        embedded_features = list(to_numpy(encoder(images)))
        labels = list(to_numpy(labels))
        break
    return embedded_features, labels


def visualise_clusters(args, model_path, vector_dim, num_of_images=1000, label_map=None, doc_file=None):
    ''' visulize and save clusters (created on feature vectors)
    '''
    if not os.path.exists(model_path):
        print('%s path does not exists!'%model_path)
        return
    embedded_features, labels = get_features(args, model_path, vector_dim, num_of_images=num_of_images)
    plot_tsne_clusters(embedded_features, labels, num_of_images, vector_dim=vector_dim, label_map=label_map,doc_file=doc_file)


def save_reconstructed_image(args, model_path, num_of_images=2, doc_file=None, show=False):
    ''' Recontruct image using decoder and save origonal and reconstructed image
    '''
    filters = [int(args.width_scale_factor*i) for i in [32, 64, 128, 256]]
    if (model_path.find('unet_based') != -1):
        model = Unet_based.End_to_End_model(3, 10, filters=filters, dense_size = 512, drop_rate=0.5)
    else:
        model = Mobilenetv2_based.End_to_End_model(n_classes=10, dense_size=1280)
    model.load_state_dict(torch.load(model_path), strict=False)
    model.eval()
    print('%s weights loaded successfully!'%model_path)
    testLoader = loadCIFAR10_testset(batch_size = num_of_images)
    print('%s images to reconstruct'%num_of_images)
    for images, labels in testLoader:
        output = model(images)
        reconstructed_images, preds = list([to_numpy(output[0]), to_numpy(output[1])])
        break
    reconstructed_images = reconstructed_images.transpose((0,2,3,1))
    images = to_numpy(images).transpose((0,2,3,1))
    if doc_file:
        document = Document(doc_file)
        sep = '************************************************************'
        p = document.add_paragraph(sep)
        r = p.add_run('\nReconstrcted image samples:')
        r.bold = True
        for i in range(num_of_images):
            img = np.hstack([images[i], reconstructed_images[i]])
            tmp_path = 'temp.jpg'
            matplotlib.image.imsave(tmp_path, img)
            document.add_picture(tmp_path, width=Inches(3))
            p = document.add_paragraph('Actual class=%s, Predicted class = %s'%(to_numpy(labels[i]),np.argmax(preds[i])))
        document.save(doc_file)
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    if show:
        plt.imshow(img), plt.show()
