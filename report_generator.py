import os
import pandas as pd
import seaborn as sns
import numpy as np
from collections import OrderedDict
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import warnings
warnings.filterwarnings('ignore')
from visualization_utils import *


def plot_cm(cm, labels=None, with_percent=False, show=False, doc_file=None):
    '''plot confusion matrix
    params:
        cm: python list or numpy array, confusion matrix
        with_percent: bool, visualize matrix with percent
        show: bool, visualize image
        doc_file: str, to save plot in doc file
    '''
    cm = np.array(cm) if isinstance(cm,list) else cm
    cm_sum = np.sum(cm, axis=1, keepdims=True)
    cm_perc = cm / cm_sum.astype(float) * 100
    annot = np.empty_like(cm).astype(str)
    nrows, ncols = cm.shape
    for i in range(nrows):
        for j in range(ncols):
            c, p = cm[i, j], cm_perc[i, j]
            if i == j:
                annot[i, j] = '%.1f%%\n%d'%(p, c)
            elif c == 0:
                annot[i, j] = ''
            else:
                if with_percent:
                    annot[i, j] = '%.1f%%\n%d'%(p, c)
                else:
                    annot[i, j] = '%d'%(c)
    cm = pd.DataFrame(cm)
    cm.index.name = 'Actual'
    cm.columns.name = 'Predicted'
    if labels:
        cm.columns = labels
        cm.index = labels
    fig, ax = plt.subplots(figsize=(7.5,4.35))
    _ = sns.heatmap(cm, annot=annot, fmt='', ax=ax)
    if doc_file:
        tmp_path = 'temp.jpg'
        plt.savefig(tmp_path,bbox_inches='tight')
        document = Document(doc_file)
        p = document.add_paragraph()
        r = p.add_run('Confusion matrix:')
        r.bold = True
        document.add_picture(tmp_path,width=Inches(5))
        document.save(doc_file)
        if os.path.exists(tmp_path):
            os.remove(tmp_path)
    if show:
        plt.show()

def get_evaluation_data(testfile_path):
    '''Rread result test file and return dataset information and evaluation information in text and list format
    '''
    count = 0
    dataset_report, cls_report, cm_report = '', [], []
    in_cls_report, in_cm_report = False, False
    labels = []
    with open(testfile_path,'r') as file:
        for line in file:
            if line.strip()=='':
                continue
            if 0<count<6:
                dataset_report += line.strip()+'\n'

            if line.strip()=='Classwise report:':
                in_cls_report = True
                continue
            if line.strip().split(' ')[0]=='accuracy':
                acc = [i for i in line.strip().split(' ')  if i!='']
                acc = [acc[0],'','','',acc[-2],acc[-1]]
                cls_report.append(acc)
                continue
            if in_cls_report:
                cls_report.append([i for i in line.strip().split(' ') if i!=''])
            if line.strip().split(' ')[0]=='weighted':
                in_cls_report = False

            if line.strip()=='Confusion Matrix:':
                in_cm_report = True
                continue
            if in_cm_report and (line.strip()!=''):
                cm_report.append(np.array([int(i) for i in line.strip().replace('[','').replace(']','').split(' ') if i!='']))
            count+=1
    cls_report = [OrderedDict(zip(['class','label','precision', 'recall', 'f1-score', 'support'],i)) for i in cls_report[1:]]
    labels = [i['class'] for i in cls_report if i['class'] not in ['accuracy','macro','weighted']]
    return dataset_report, cls_report, cm_report, labels

def generate_experiment_report(args, testfile_path, modelfile_path, reportfile_path):
    '''generate docx file and save evaluation results extrated from testfile_path 
        and save encoded feature clusters and 2 sample reconstructed image
    '''
    dataset_report, cls_report, cm_report, labels = get_evaluation_data(testfile_path)
    if os.path.exists(reportfile_path):
        document = Document()
    else:
        document = Document()

    sep = '************************************************************'
    training_type = 'End to end training' if args.is_end_to_end_training else 'Separate training' if args.is_separate_training else 'Classification training'
    exp_info = f'Training: {training_type}\nArchitecture: {args.arch}\nAE loss: {args.ae_loss}\nClassifier loss: {args.classifier_loss}\nOptimizer: {args.optimizer}\nData sampling strategy: {args.data_sampling}\nLearning rate: {args.learning_rate}\nBatch size: {args.batch_size}\nEpochs: {args.num_epochs}'
    
    p = document.add_paragraph()
    r = p.add_run('Training Hyperparameters')
    r.bold = True
    p = document.add_paragraph('%s'%(exp_info))
    r = p.add_run(sep)
    p = document.add_paragraph()
    r = p.add_run('Dataset')
    r.bold = True
    p = document.add_paragraph(dataset_report)
    p.style = 'List Bullet'
    r = p.add_run(sep)
    document.save(reportfile_path)

    df = pd.DataFrame.from_dict(cls_report)
    df.columns = ['class','label','precision', 'recall', 'f1-score', 'support']
    p = document.add_paragraph()
    r = p.add_run('Classification report:')
    r.bold = True
    t = document.add_table(df.shape[0]+1, df.shape[1])
    t.style = 'LightList'
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    p = document.add_paragraph(sep)
    document.save(reportfile_path)

    plot_cm(cm_report, labels = labels, doc_file=reportfile_path)
    visualise_clusters(args, modelfile_path, vector_dim=320, num_of_images=1000, doc_file=reportfile_path)
    save_reconstructed_image(args, modelfile_path, num_of_images=2, doc_file=reportfile_path)


